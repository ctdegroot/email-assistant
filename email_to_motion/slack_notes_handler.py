"""
slack_notes_handler.py — Process messages in the #notes-inbox Slack channel.

When an email is forwarded to the channel via Slack's email integration, this
module:
  1. Extracts the message text (email body) and any file attachments.
  2. Downloads and parses PDF and Word (.docx) attachments for their text.
  3. Computes a SHA-256 source hash of the email content for deduplication.
  4. Calls note_generator to produce structured Obsidian-flavoured Markdown.
  5. Writes the .md file to NOTES_OUTPUT_PATH (dedup: overwrites if re-send,
     creates a (2) file if genuinely different email with same subject/date).
  6. Optionally pushes the note to the Obsidian vault via Git
     (when OBSIDIAN_DELIVERY=git).
  7. Posts a confirmation back to the channel — "saved", "updated", or
     "unchanged" — and marks the original message with ✅.

Error handling:
  On Claude or file-write failure: posts a ⚠️ warning to the channel and
  sends a DM to ALLOWED_SLACK_USER_ID (if configured) with details.
  Transient Claude API errors (rate limits, timeouts) are retried
  automatically with exponential back-off via call_with_retries().

Supported attachment types:
  - PDF  (.pdf)  — extracted with pdfplumber
  - Word (.docx) — extracted with python-docx
  Other file types are listed in the note frontmatter but not parsed for text.
"""

import io
import json
import logging
import re
from datetime import datetime
from pathlib import Path

import requests

from . import config
from . import note_generator
from . import vault_writer
from .slack_helpers import get_channel_id, get_unprocessed_messages, mark_processed

log = logging.getLogger(__name__)


# ── Channel ID resolution ─────────────────────────────────────────────────────
# Incoming Slack events carry a channel ID, not a name. We resolve the configured
# channel name once at startup and cache it here for fast comparison in _dispatch.

_channel_id: str = ""

# ── In-flight deduplication ───────────────────────────────────────────────────
# Tracks message timestamps currently being processed by the socket listener.
# The startup sweep checks this set to avoid re-processing a message that the
# socket thread has already started on (but hasn't yet added the ✅ reaction).
_processing_ts: set[str] = set()


def init(channel_name: str):
    """
    Resolve SLACK_NOTES_CHANNEL (a name) to its Slack channel ID and cache it.
    Called once from socket_listener.start() after the Slack client is ready.
    Does nothing if channel_name is empty (feature disabled).
    """
    global _channel_id
    if not channel_name:
        return

    try:
        _channel_id = get_channel_id(channel_name.lstrip("#"))
    except ValueError as e:
        print(f"⚠️  {e} — note processing disabled.")
    except Exception as e:
        print(f"⚠️  Could not resolve notes channel '{channel_name}': {e}")


# ── Attachment extraction ─────────────────────────────────────────────────────

def _download_slack_file(url: str) -> bytes:
    """Download a private Slack file, authenticating with the bot token."""
    resp = requests.get(
        url,
        headers={"Authorization": f"Bearer {config.SLACK_BOT_TOKEN}"},
        timeout=30,
    )
    resp.raise_for_status()
    return resp.content


def _extract_pdf(data: bytes) -> str:
    """Extract plain text from a PDF using pdfplumber."""
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return "\n\n".join(p.strip() for p in pages if p.strip())
    except Exception as e:
        return f"(PDF extraction failed: {e})"


def _extract_docx(data: bytes) -> str:
    """Extract plain text from a Word .docx file using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(data))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append("  |  ".join(cells))
        return "\n".join(paragraphs)
    except Exception as e:
        return f"(Word extraction failed: {e})"


# Map both MIME type and Slack's short filetype string to an extractor function
_BY_MIME = {
    "application/pdf": _extract_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_docx,
}
_BY_FILETYPE = {
    "pdf":  _extract_pdf,
    "docx": _extract_docx,
}

# Reverse map: MIME → short filetype string (used when normalising email attachments)
_MIME_TO_FT: dict[str, str] = {v: k for k, v in {
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}.items()}


def _mime_to_filetype(mime: str) -> str:
    """Return a short filetype string for a MIME type, or derive one from the subtype."""
    if mime in _MIME_TO_FT:
        return _MIME_TO_FT[mime]
    # Fallback: take the part after '/' and strip vendor prefixes
    # e.g. "image/jpeg" → "jpeg", "application/vnd.ms-excel" → "ms-excel"
    _, _, subtype = mime.partition("/")
    subtype = subtype.split(";")[0].strip()
    for prefix in ("vnd.", "x-"):
        if subtype.startswith(prefix):
            subtype = subtype[len(prefix):]
    return subtype or ""


# ── Forwarded email unwrapping ────────────────────────────────────────────────

# Matches common forwarded-message divider lines, e.g.:
#   "---------- Forwarded message ---------"
#   "-----Original Message-----"
#   "Begin forwarded message:"
_FWD_MARKER_RE = re.compile(
    r'(-{3,}\s*(forwarded message|original message)\s*-{3,}'
    r'|begin forwarded message\s*:)',
    re.IGNORECASE,
)

# Strips one or more Fwd:/FW: prefixes from a subject line
_FWD_SUBJECT_RE = re.compile(r'^(fwd?\s*:\s*)+', re.IGNORECASE)


def _unwrap_forward(subject: str, body: str) -> tuple[str, str | None]:
    """
    If this looks like a forwarded email, return the original subject and sender.

    Strategy:
      1. Strip Fwd:/FW: prefix(es) from the subject line.
      2. Scan the body for a forwarded-message block (common divider patterns).
         If found, search that block for a "From:" line containing the original
         sender. If no divider is found, fall back to the first "From:" line in
         the entire body.

    Returns (clean_subject, original_sender | None).
    original_sender is None if nothing useful was found in the body.
    """
    clean_subject = _FWD_SUBJECT_RE.sub("", subject).strip()

    # Find the section of the body that contains the forwarded headers
    m = _FWD_MARKER_RE.search(body)
    search_body = body[m.end():] if m else body

    # Extract the original From: line
    from_match = re.search(r'(?im)^from:\s*(.+)', search_body)
    original_sender = from_match.group(1).strip() if from_match else None

    return clean_subject, original_sender


# ── Email content extraction ──────────────────────────────────────────────────

def _extract_email_content(files: list[dict]) -> tuple[str, str, str, list[dict]]:
    """
    Extract email metadata and body from Slack's email integration file structure.

    When an email is forwarded to a Slack channel via the channel email address,
    Slack posts a message where the email itself is stored as a file with
    filetype='email', carrying structured fields (subject, from, plain_text).
    Any attachments from the original email appear as separate file entries.

    Returns:
        subject        — email subject line, or 'Untitled'
        sender         — sender display string, or 'Unknown'
        body           — plain text body of the email
        attachment_files — list of non-email file dicts (PDFs, Word docs, etc.)
    """
    subject          = "Untitled"
    sender           = "Unknown"
    body             = ""
    attachment_files = []

    for f in files:
        if f.get("filetype") == "email":
            subject = f.get("subject") or "Untitled"
            frm     = (f.get("from") or [{}])
            sender  = frm[0].get("original") or frm[0].get("address") or "Unknown"
            body    = f.get("plain_text") or ""

            # Attachments from the original email can appear in two places
            # depending on the Slack API version / email type:
            #
            # 1. f["files"]       — regular Slack file objects (same shape as
            #                       top-level event files; used in some integrations)
            # 2. f["attachments"] — Slack's email-attachment objects, which use
            #                       slightly different field names and must be
            #                       normalised before the downloader can use them.
            for embedded in (f.get("files") or []):
                attachment_files.append(embedded)

            for att in (f.get("attachments") or []):
                # Email attachment objects have url_private / url_private_download
                # directly, but use 'filename' (not 'name') for the file name and
                # may omit 'filetype' (using 'mimetype' instead).
                # Field priority confirmed from live Slack payloads:
                #   filename > name > title
                name = (att.get("filename") or att.get("name") or att.get("title") or "attachment")
                mime = att.get("mimetype") or ""
                # Derive a short filetype string from the MIME type if absent
                ft_short = att.get("filetype") or _mime_to_filetype(mime)
                dl_url   = (att.get("url_private_download")
                            or att.get("url_private")
                            or att.get("url")
                            or "")
                if dl_url or mime:
                    attachment_files.append({
                        "name":                 name,
                        "mimetype":             mime,
                        "filetype":             ft_short,
                        "url_private_download": dl_url,
                        "url_private":          dl_url,
                    })
        else:
            attachment_files.append(f)

    # If this is a forwarded email, replace the forwarder's details with the
    # original sender and strip the Fwd:/FW: prefix from the subject.
    if _FWD_SUBJECT_RE.match(subject) or _FWD_MARKER_RE.search(body):
        clean_subject, original_sender = _unwrap_forward(subject, body)
        subject = clean_subject or subject
        if original_sender:
            sender = original_sender

    return subject, sender, body, attachment_files


# ── Startup sweep ─────────────────────────────────────────────────────────────

def process_unprocessed_notes():
    """
    On startup, fetch any unprocessed messages from the notes-inbox channel
    and run them through the note pipeline. Mirrors the behaviour of
    process_channel() and process_calendar_channel() for the other inboxes.

    A message is considered unprocessed if it has no ✅ reaction.
    Does nothing if the notes channel is not configured.
    """
    if not _channel_id:
        return

    print(f"Checking #{config.SLACK_NOTES_CHANNEL}…")
    msgs = get_unprocessed_messages(_channel_id)
    if not msgs:
        print("  No unprocessed notes.")
        return

    for msg in msgs:
        # Skip any message already being handled by the socket listener thread.
        # Without this guard, a message that arrived while the service is running
        # would be processed twice: once in real time via the socket, and again
        # here before the socket thread has had time to add the ✅ reaction.
        if msg.get("ts") in _processing_ts:
            log.info("notes sweep: skipping ts=%s — already in flight", msg.get("ts"))
            continue
        # conversations_history messages don't carry a 'channel' field —
        # inject it so process_message() can post confirmations and mark ✅
        process_message({**msg, "channel": _channel_id})


# ── Main entry point ──────────────────────────────────────────────────────────

def process_message(event: dict):
    """
    Process one message event from the notes-inbox channel.

    Slack's email integration stores forwarded emails as a file with
    filetype='email' rather than in event['text']. This function extracts
    content from that structure, then processes any PDF or Word attachments.

    Runs in a background thread — never called from _dispatch directly.
    """
    # Guard: ignore our own bot's messages to avoid feedback loops
    if event.get("bot_id") == config.OWN_BOT_ID:
        return

    # Ignore structural non-message events
    subtype = event.get("subtype", "")
    if subtype in ("message_changed", "message_deleted", "channel_join", "channel_leave"):
        return

    # Register this message as in-flight so the startup sweep doesn't also
    # process it. Always deregister on exit (success or exception) so the set
    # never grows stale.
    ts = event.get("ts", "")
    if ts in _processing_ts:
        log.info("notes: skipping duplicate event ts=%s", ts)
        return
    if ts:
        _processing_ts.add(ts)

    try:
        _process_message_inner(event)
    finally:
        _processing_ts.discard(ts)


# ── Inner-body extraction for content hashing ────────────────────────────────

# Matches common email metadata header lines that appear inside a forwarded block,
# e.g. "From: ...", "Date: ...", "Subject: ...", "To: ...", "Cc: ...", "Reply-To: ..."
_EMAIL_HEADER_RE = re.compile(r'^[A-Za-z][A-Za-z-]+:\s', re.MULTILINE)


def _extract_inner_body(body: str) -> str:
    """
    Return only the *original* email content from a (potentially forwarded) body.

    When an email is forwarded, the body typically looks like:

        [User's forwarding note / signature]
        ---------- Forwarded message ---------
        From: orig@example.com
        Date: ...
        Subject: ...
        To: ...

        Actual original email content here.

    This function:
      1. Finds the forwarding divider (same regex as _FWD_MARKER_RE).
      2. Skips any email-header lines that immediately follow (From:, Date: …).
      3. Returns the trimmed inner content.

    If no forwarding divider is found the body is returned unchanged — hashing a
    non-forwarded email body should work as-is.

    Used by compute_source_hash so that differences in the forwarding wrapper
    (changed signature, extra "FYI" note, etc.) don't affect the duplicate-
    detection hash.
    """
    m = _FWD_MARKER_RE.search(body)
    if not m:
        return body.strip()

    # Everything after the divider, split into lines
    lines = body[m.end():].splitlines()

    # Skip leading blank lines and metadata header lines
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if stripped == "" or _EMAIL_HEADER_RE.match(lines[i]):
            i += 1
        else:
            break

    return "\n".join(lines[i:]).strip()


# ── Owner DM alerting ─────────────────────────────────────────────────────────

def _dm_owner_on_failure(subject: str, channel_id: str, error: Exception):
    """
    Send a direct message to the bot owner (ALLOWED_SLACK_USER_ID) when note
    generation fails.  This provides a persistent, out-of-channel record of
    failures so nothing is silently lost.

    Only fires when ALLOWED_SLACK_USER_ID is configured.  Swallows any
    exceptions so a DM failure never masks the original error.
    """
    owner = config.ALLOWED_SLACK_USER_ID
    if not owner:
        return
    try:
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
        lines = [
            "⚠️ *Note generation failed*",
            f"• Subject: _{subject}_",
            f"• Channel: <#{channel_id}>",
            f"• Error: `{type(error).__name__}: {error}`",
            f"• Time: {now_str}",
        ]
        config.slack.chat_postMessage(channel=owner, text="\n".join(lines))
    except Exception:
        pass   # non-fatal: DM failure must never obscure the original error



# Maps write_note() status strings to the Slack confirmation label.
_WRITE_STATUS_LABEL: dict[str, str] = {
    "saved":     "📝 Note saved",
    "updated":   "📝 Note updated",
    "unchanged": "📝 Note unchanged",
}


def _process_message_inner(event: dict):
    """Inner implementation — called only after the dedup guard in process_message()."""
    channel_id = event.get("channel", "")
    files      = event.get("files") or []

    # ── Log a brief summary; set LOG_LEVEL=DEBUG for the full file-object dump ──
    if files:
        for i, f in enumerate(files):
            ft = f.get("filetype", "?")
            atts = f.get("attachments") or [] if ft == "email" else []
            att_names = [a.get("filename") or a.get("name") or "?" for a in atts]
            log.info("notes file[%d]: filetype=%r  attachments=%s",
                     i, ft, att_names or "(none)")
            if ft == "email":
                safe = {k: v for k, v in f.items() if k not in ("plain_text", "preview_plain_text")}
                log.debug("  email file[%d] full object: %s", i, json.dumps(safe, default=str))
    else:
        log.info("notes event: no files attached")

    # Extract email content from Slack's email file structure
    subject, sender, body, attachment_files = _extract_email_content(files)

    # Fall back to raw message text if no email file was present
    # (e.g. someone typed directly into the channel for testing)
    if not body:
        body = (event.get("text") or "").strip()

    date_str = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── Download and extract attachment text ──────────────────────────────────
    attachment_names: list[str] = []
    attachment_texts: dict[str, str] = {}

    for f in attachment_files:
        fname    = f.get("name") or "attachment"
        mimetype = f.get("mimetype") or ""
        filetype = (f.get("filetype") or "").lower()
        dl_url   = f.get("url_private_download") or f.get("url_private") or ""

        attachment_names.append(fname)

        extractor = _BY_MIME.get(mimetype) or _BY_FILETYPE.get(filetype)
        if extractor and dl_url:
            try:
                raw = _download_slack_file(dl_url)
                attachment_texts[fname] = extractor(raw)
            except Exception as e:
                attachment_texts[fname] = f"(download failed: {e})"
        # Unsupported types are listed in frontmatter but not parsed for text

    # ── Fingerprint the source email so we can detect re-sends later ──────────
    # Hash the *inner* body only (strips the user's forwarding wrapper and
    # signature so that re-forwarding with a different note or signature still
    # produces the same hash as the original send).
    source_hash = note_generator.compute_source_hash(
        subject=subject,
        sender=sender,
        body=_extract_inner_body(body),
        attachment_names=attachment_names,
    )

    # ── Generate the Markdown note via Claude ─────────────────────────────────
    try:
        markdown = note_generator.generate_note(
            subject=subject,
            sender=sender,
            date=date_str,
            body=body,
            attachment_names=attachment_names,
            attachment_texts=attachment_texts,
        )
    except Exception as e:
        log.error("notes: note generation failed for %r: %s", subject, e)
        config.slack.chat_postMessage(
            channel=channel_id,
            text=f"⚠️ Could not generate note for *{subject}*: {e}",
        )
        _dm_owner_on_failure(subject, channel_id, e)
        return

    # ── Write the .md file locally ────────────────────────────────────────────
    output_dir = Path(config.NOTES_OUTPUT_PATH).expanduser().resolve()
    try:
        written_path, write_status = note_generator.write_note(
            markdown, subject, output_dir, source_hash=source_hash,
        )
    except Exception as e:
        log.error("notes: note file write failed for %r: %s", subject, e)
        config.slack.chat_postMessage(
            channel=channel_id,
            text=f"⚠️ Could not write note file for *{subject}*: {e}",
        )
        _dm_owner_on_failure(subject, channel_id, e)
        return

    # ── Optionally push to Obsidian vault via Git ─────────────────────────────
    vault_path: Path | None = None
    vault_error: str        = ""

    if config.OBSIDIAN_DELIVERY == "git":
        try:
            vault_path = vault_writer.push_to_vault(markdown, written_path.name)
        except Exception as e:
            vault_error = str(e)

    # Mark the original message as processed so it's skipped on the next sweep
    ts = event.get("ts", "")
    if ts:
        mark_processed(channel_id, ts)

    # ── Confirm in Slack ──────────────────────────────────────────────────────
    att_line = ""
    if attachment_names:
        parsed   = [n for n in attachment_names if n in attachment_texts]
        unparsed = [n for n in attachment_names if n not in attachment_texts]
        parts    = []
        if parsed:
            parts.append(f"parsed: {', '.join(parsed)}")
        if unparsed:
            parts.append(f"listed only: {', '.join(unparsed)}")
        att_line = f"\n  Attachments — {'; '.join(parts)}"

    if vault_path:
        delivery_line = f"\n  📚 Pushed to vault: `{vault_path}`"
    elif vault_error:
        delivery_line = f"\n  ⚠️ Vault push failed: {vault_error}"
    else:
        delivery_line = f"\n  `{written_path}`"

    status_label = _WRITE_STATUS_LABEL.get(write_status, "📝 Note saved")

    config.slack.chat_postMessage(
        channel=channel_id,
        text=(
            f"{status_label}: *{written_path.name}*"
            f"{att_line}"
            f"{delivery_line}"
        ),
    )
